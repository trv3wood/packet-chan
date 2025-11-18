"""Simple ingestion helpers: extract text and chunk it.

Supports: .txt, .md, .pdf, .docx
Extend this module to integrate Markitdown or other converters for PPTX/images.
"""
from typing import List
import os
import tempfile
import warnings
import docx
import logging

# 配置日志
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# 尝试使用 pypdf（PyPDF2 的升级版，对中文支持更好）
try:
    from pypdf import PdfReader
    USE_PYPDF = True
except ImportError:
    try:
        from PyPDF2 import PdfReader
        USE_PYPDF = False
        warnings.warn("建议使用 pypdf 替代 PyPDF2 以获得更好的中文支持。运行: pip install pypdf")
    except ImportError:
        PdfReader = None
        warnings.warn("未安装 PDF 处理库。运行: pip install pypdf")

# Try to import the LangChain text splitter; provide a lightweight fallback if unavailable.
try:
    from langchain_text_splitters import RecursiveCharacterTextSplitter
except Exception:
    warnings.warn(
        "langchain not available or RecursiveCharacterTextSplitter import failed; "
        "using a simple fallback splitter. For better splitting, install langchain."
    )

    class _SimpleRecursiveCharacterTextSplitter:
        """A very small fallback to mimic RecursiveCharacterTextSplitter.split_text."""

        def __init__(self, chunk_size=1000, chunk_overlap=200, length_function=len, separators=None):
            self.chunk_size = chunk_size
            self.chunk_overlap = chunk_overlap
            self.length_function = length_function
            self.separators = separators or ["\n\n", "\n", ". ", "! ", "? ", " ", ""]

        def split_text(self, text: str) -> List[str]:
            if not text:
                return []

            # Try to split by the first available separator to produce more natural chunks.
            for sep in self.separators:
                if sep and sep in text:
                    parts = text.split(sep)
                    chunks: List[str] = []
                    for part in parts:
                        part = part.strip()
                        if not part:
                            continue
                        # Ensure parts that are too long are further split with overlap.
                        i = 0
                        while i < len(part):
                            end = min(i + self.chunk_size, len(part))
                            chunks.append(part[i:end].strip())
                            if end >= len(part):
                                break
                            i += max(1, self.chunk_size - self.chunk_overlap)
                    return chunks

            # Fallback sliding-window splitting if no separator matched.
            chunks = []
            start = 0
            text_len = len(text)
            while start < text_len:
                end = min(start + self.chunk_size, text_len)
                chunks.append(text[start:end].strip())
                if end >= text_len:
                    break
                start = max(0, end - self.chunk_overlap)
            return chunks

    # expose the fallback under the expected name
    RecursiveCharacterTextSplitter = _SimpleRecursiveCharacterTextSplitter



def extract_text_from_pdf(path: str) -> str:
    """从 PDF 文件提取文本，支持多种提取方法和编码"""
    logger.info("开始提取PDF文本")
    # 首先尝试 pypdf
    try:
        logger.info("尝试使用PyPDF提取")
        return _extract_with_pypdf(path)
    except Exception as e:
        logger.warning(f"PyPDF提取失败: {e}, 尝试pdfplumber")
        # 如果 pypdf 失败，尝试 pdfplumber
        try:
            import pdfplumber
            logger.info("尝试使用pdfplumber提取")
            return _extract_with_pdfplumber(path)
        except ImportError:
            logger.error("pdfplumber未安装")
            raise ValueError(
                f"PDF 提取失败: {e}\n"
                f"建议安装 pdfplumber 以获得更好的支持: pip install pdfplumber"
            )
        except Exception as e2:
            logger.error(f"pdfplumber提取失败: {e2}")
            raise ValueError(f"PDF 提取失败（pypdf 和 pdfplumber 都失败）: {e}, {e2}")


def _extract_with_pypdf(path: str) -> str:
    """使用 pypdf 提取文本"""
    logger.info("创建PdfReader对象")
    text = []
    with warnings.catch_warnings():
        warnings.filterwarnings("ignore", category=UserWarning)
        warnings.filterwarnings("ignore", message=".*encoding.*")
        
        reader = PdfReader(path, strict=False)
        total_pages = len(reader.pages)
    
    logger.info(f"PDF总页数: {total_pages}")
    for page_num, page in enumerate(reader.pages, 1):
        logger.info(f"提取第 {page_num}/{total_pages} 页")
        try:
            page_text = page.extract_text()
            if page_text and len(page_text.strip()) > 0:
                text.append(page_text)
            else:
                logger.warning(f"第 {page_num} 页：无法提取文本")
                text.append(f"[第 {page_num} 页：无法提取文本]")
        except Exception as e:
            logger.error(f"第 {page_num} 页：提取错误 - {str(e)}")
            text.append(f"[第 {page_num} 页：提取错误 - {str(e)}]")
    
    result = "\n\n".join(text)
    actual_text = result.replace("[第", "").replace("页：无法提取文本]", "").replace("页：提取错误", "").strip()
    
    logger.info(f"PyPDF提取完成，共提取文本长度: {len(actual_text)} 字符")
    
    if not actual_text or len(actual_text) < 10:
        raise ValueError(f"PDF 文本提取失败：提取的文本太短或为空")
    
    return result


def _extract_with_pdfplumber(path: str) -> str:
    """使用 pdfplumber 提取文本（回退方法）"""
    import pdfplumber
    logger.info("使用pdfplumber打开PDF")
    text = []
    
    with pdfplumber.open(path) as pdf:
        total_pages = len(pdf.pages)
        logger.info(f"PDF总页数: {total_pages}")
        for page_num, page in enumerate(pdf.pages, 1):
            logger.info(f"提取第 {page_num}/{total_pages} 页")
            try:
                page_text = page.extract_text()
                if page_text:
                    text.append(page_text)
                else:
                    logger.warning(f"第 {page_num} 页：无法提取文本")
                    text.append(f"[第 {page_num} 页：无法提取文本]")
            except Exception as e:
                logger.error(f"第 {page_num} 页：提取错误 - {str(e)}")
                text.append(f"[第 {page_num} 页：提取错误 - {str(e)}]")
    
    result = "\n\n".join(text)
    actual_text = result.replace("[第", "").replace("页：无法提取文本]", "").replace("页：提取错误", "").strip()
    
    logger.info(f"pdfplumber提取完成，共提取文本长度: {len(actual_text)} 字符")
    
    if not actual_text or len(actual_text) < 10:
        raise ValueError(f"PDF 文本提取失败：提取的文本太短或为空")
    
    return result


def extract_text_from_docx(path: str) -> str:
    logger.info("打开Word文档")
    doc = docx.Document(path)
    paragraphs = []
    
    total_paragraphs = len(doc.paragraphs)
    for para_num, p in enumerate(doc.paragraphs):
        if para_num % 10 == 0:  # 每10段记录一次日志
            logger.info(f"处理第 {para_num + 1}/{total_paragraphs} 段")
        paragraphs.append(p.text)
    
    logger.info(f"Word文档提取完成，共处理段落数: {total_paragraphs}")
    return "\n\n".join(paragraphs)


def extract_text_from_file(path: str) -> str:
    logger.info(f"开始处理文件: {path}")
    _, ext = os.path.splitext(path.lower())
    logger.info(f"文件扩展名: {ext}")
    
    if ext in (".txt", ".md"):
        logger.info(f"文件类型为{ext.upper()}文本文件")
        try:
            with open(path, "r", encoding="utf-8", errors="ignore") as f:
                content = f.read()
                logger.info(f"文本文件读取完成，内容长度: {len(content)} 字符")
                return content
        except Exception as e:
            logger.error(f"读取文本文件失败: {e}")
            raise
    if ext == ".pdf":
        logger.info("文件类型为PDF")
        return extract_text_from_pdf(path)
    if ext == ".docx":
        logger.info("文件类型为Word文档")
        return extract_text_from_docx(path)
    # TODO: add pptx, images (OCR), other converters (Markitdown)
    logger.error(f"不支持的文件类型: {ext}")
    raise ValueError(f"Unsupported file type: {ext}")


def chunk_text(text: str, chunk_size: int = 1000, overlap: int = 200) -> List[str]:
    """LangChain RecursiveCharacterTextSplitter for intelligent text chunking.

    Returns list of text chunks.
    """
    logger.info(f"开始文本分块，总长度: {len(text)} 字符")
    if not text:
        logger.info("文本为空，返回空列表")
        return []
    
    # Create text splitter with the same parameters as the original function
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=overlap,
        length_function=len,
        separators=["\n\n", "\n", ". ", "! ", "? ", " ", ""]
    )
    
    # Split the text into chunks
    chunks = text_splitter.split_text(text)
    logger.info(f"文本分块完成，生成 {len(chunks)} 个块")
    return chunks


def save_uploaded_file(uploaded_file, dst_dir: str) -> str:
    """Given a Streamlit UploadedFile-like object, write to a temp file and return path."""
    logger.info(f"保存上传文件: {uploaded_file.name}")
    logger.info(f"目标目录: {dst_dir}")
    
    os.makedirs(dst_dir, exist_ok=True)
    
    # 获取原始文件的扩展名
    original_name = uploaded_file.name
    _, ext = os.path.splitext(original_name)
    logger.info(f"文件扩展名: {ext}")
    
    # 保留原始扩展名，这样 extract_text_from_file 才能正确识别文件类型
    fd, path = tempfile.mkstemp(prefix="upload_", dir=dst_dir, suffix=ext)
    os.close(fd)
    
    logger.info(f"准备写入临时文件: {path}")
    with open(path, "wb") as f:
        f.write(uploaded_file.getbuffer())
    
    logger.info(f"文件保存成功，路径: {path}")
    return path
